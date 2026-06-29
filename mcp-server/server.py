import os

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from starlette.concurrency import run_in_threadpool
from pydantic import BaseModel
from typing import Any, Dict, List
from io import BytesIO

from docx import Document
from pypdf import PdfReader

from tools import (
    architecture_diagram_generator,
    assistant_agent,
    aws_cost_analyzer,
    backend_developer,
    bedrock_text_generator,
    create_ppt,
    devops_engineer,
    frontend_developer,
    quality_engineer,
    review_python_code,
    sql_generator,
    summarize_document,
)

app = FastAPI(title="MCP Server POC")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class ToolRequest(BaseModel):
    tool: str
    input: str


class ToolInfo(BaseModel):
    name: str
    description: str
    inputs: List[str]
    outputs: List[str]


ALLOWED_SUMMARY_EXTENSIONS = {".txt", ".pdf", ".docx"}
MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_EXTRACTED_CHARS = int(os.getenv("MAX_EXTRACTED_CHARS", "1000000"))


def get_file_extension(filename: str) -> str:
    if "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[1].lower()


def extract_text_from_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(encoding)[:MAX_EXTRACTED_CHARS]
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=400, detail="Unable to decode text file.")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(file_bytes))
        page_text = []
        for page in reader.pages:
            extracted = page.extract_text() or ""
            if extracted.strip():
                page_text.append(extracted.strip())
            if sum(len(text) for text in page_text) >= MAX_EXTRACTED_CHARS:
                break
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Unable to read PDF file: {exc}") from exc
    return "\n\n".join(page_text)[:MAX_EXTRACTED_CHARS]


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Unable to read DOCX file: {exc}") from exc

    paragraphs = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)
        if sum(len(value) for value in paragraphs) >= MAX_EXTRACTED_CHARS:
            break
    table_cells = [
        cell.text.strip()
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    ]
    combined_text = "\n".join(paragraphs + table_cells)
    return combined_text[:MAX_EXTRACTED_CHARS]


def extract_text_from_upload(filename: str, file_bytes: bytes) -> str:
    extension = get_file_extension(filename)
    if extension not in ALLOWED_SUMMARY_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Upload a .txt, .pdf, or .docx file.",
        )
    if len(file_bytes) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=400, detail="File is too large. Maximum size is 10 MB.")

    if extension == ".txt":
        return extract_text_from_txt(file_bytes)
    if extension == ".pdf":
        return extract_text_from_pdf(file_bytes)
    return extract_text_from_docx(file_bytes)


TOOL_REGISTRY = {
    "review_python_code": {
        "description": "Review Python code for style, logic, best practices, and potential bugs.",
        "inputs": ["code"],
        "outputs": ["summary", "issues"],
        "function": review_python_code,
    },
    "create_ppt": {
        "description": "Generate a PowerPoint slide outline from text input.",
        "inputs": ["text"],
        "outputs": ["summary", "slides"],
        "function": create_ppt,
    },
    "summarize_document": {
        "description": "Summarize documents and long-form text with key points and action items.",
        "inputs": ["text"],
        "outputs": ["summary", "key_points", "recommendation"],
        "function": summarize_document,
    },
    "aws_cost_analyzer": {
        "description": "Analyze AWS cost data and provide optimization insights.",
        "inputs": ["data"],
        "outputs": ["summary", "insights"],
        "function": aws_cost_analyzer,
    },
    "sql_generator": {
        "description": "Generate or optimize SQL queries based on natural language prompts.",
        "inputs": ["prompt"],
        "outputs": ["summary", "query", "notes"],
        "function": sql_generator,
    },
    "architecture_diagram_generator": {
        "description": "Create architecture diagram descriptions for cloud and application designs.",
        "inputs": ["description"],
        "outputs": ["summary", "diagram_steps", "recommendation"],
        "function": architecture_diagram_generator,
    },
    "frontend_developer": {
        "description": "Review frontend code, UI/UX, accessibility, performance, and compatibility.",
        "inputs": ["input_text"],
        "outputs": ["summary", "recommendations"],
        "function": frontend_developer,
    },
    "devops_engineer": {
        "description": "Evaluate deployment, automation, CI/CD, infrastructure, and operational readiness.",
        "inputs": ["input_text"],
        "outputs": ["summary", "recommendations"],
        "function": devops_engineer,
    },
    "quality_engineer": {
        "description": "Assess test coverage, quality gates, validation, and reliability risks.",
        "inputs": ["input_text"],
        "outputs": ["summary", "observations"],
        "function": quality_engineer,
    },
    "backend_developer": {
        "description": "Review backend architecture, API design, data models, security, and scalability.",
        "inputs": ["input_text"],
        "outputs": ["summary", "suggestions"],
        "function": backend_developer,
    },
    "bedrock_text_generator": {
        "description": "Generate text using AWS Bedrock models for flexible assistant responses.",
        "inputs": ["prompt"],
        "outputs": ["tool", "model_id", "response"],
        "function": bedrock_text_generator,
    },
    "assistant_agent": {
        "description": "Run an assistant-style agent backed by Bedrock or a local fallback.",
        "inputs": ["input_text"],
        "outputs": ["summary", "assistant_response"],
        "function": assistant_agent,
    },
}


@app.get("/tools", response_model=List[ToolInfo])
def list_tools() -> List[ToolInfo]:
    return [
        ToolInfo(
            name=name,
            description=info["description"],
            inputs=info["inputs"],
            outputs=info["outputs"],
        )
        for name, info in TOOL_REGISTRY.items()
    ]


@app.post("/invoke")
def invoke_tool(request: ToolRequest) -> Dict[str, Any]:
    tool_info = TOOL_REGISTRY.get(request.tool)
    if not tool_info:
        raise HTTPException(status_code=404, detail="Tool not found")

    result = tool_info["function"](request.input)
    return {
        "tool": request.tool,
        "input": request.input,
        "result": result,
    }


@app.post("/summarize-file")
async def summarize_file(
    file: UploadFile = File(...),
    instruction: str = Form("summarize this document"),
) -> Dict[str, Any]:
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    extracted_text = await run_in_threadpool(
        extract_text_from_upload,
        file.filename or "",
        file_bytes,
    )
    if not extracted_text.strip():
        raise HTTPException(status_code=400, detail="No readable text found in uploaded file.")

    summary_input = extracted_text.strip()
    result = await run_in_threadpool(summarize_document, summary_input)
    assistant_response = (
        f"Summary for {file.filename}\n\n"
        f"{result['summary']}\n\n"
        "Key points:\n"
        + "\n".join(f"- {point}" for point in result["key_points"])
        + f"\n\nRecommendation: {result['recommendation']}"
    )
    was_truncated = len(extracted_text) >= MAX_EXTRACTED_CHARS

    return {
        "tool": "summarize_document",
        "filename": file.filename,
        "input": instruction,
        "result": {
            **result,
            "assistant_response": assistant_response,
            "extracted_characters": len(extracted_text),
            "source_file": file.filename,
            "source_type": get_file_extension(file.filename or ""),
            "truncated": was_truncated,
            "instruction": instruction,
        },
    }


@app.get("/")
def health_check() -> Dict[str, str]:
    return {"status": "ok", "message": "MCP Server is running."}
